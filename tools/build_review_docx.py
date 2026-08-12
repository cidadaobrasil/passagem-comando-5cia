from pathlib import Path

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
HTML = ROOT / "index.html"
OUT = ROOT / "revisao-conteudo-word.docx"

COLORS = {
    "navy": "17365D",
    "navy_dark": "0B2545",
    "header_fill": "E8EEF5",
    "row_alt": "F2F4F7",
    "line": "C8D0DA",
    "muted": "666666",
    "danger_fill": "FCE8E6",
    "warning_fill": "FFF4CE",
    "info_fill": "EAF2FF",
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for w in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    hdr = OxmlElement("w:tblHeader")
    hdr.set(qn("w:val"), "true")
    tr_pr.append(hdr)


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.15):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_runs(paragraph, node, bold=False):
    if isinstance(node, NavigableString):
        txt = str(node)
        if txt:
            run = paragraph.add_run(txt)
            run.bold = bold
        return
    if not isinstance(node, Tag):
        return
    next_bold = bold or node.name in {"strong", "b"}
    if node.name == "br":
        paragraph.add_run(" ")
        return
    for child in node.children:
        add_runs(paragraph, child, next_bold)


def clean_paragraph_text(paragraph):
    text = paragraph.text
    text = " ".join(text.split())
    for run in paragraph.runs:
        run.text = ""
    if text:
        paragraph.runs[0].text = text if paragraph.runs else ""


def add_html_paragraph(container, node, style=None, bullet=False):
    p = container.add_paragraph(style=style)
    if bullet:
        p.style = "List Bullet"
    add_runs(p, node)
    text = " ".join(p.text.split())
    for r in p.runs:
        r.text = ""
    if text:
        if p.runs:
            p.runs[0].text = text
        else:
            p.add_run(text)
    set_paragraph_spacing(p, after=4 if bullet else 6, line=1.18)
    return p


def style_document(doc):
    section = doc.sections[0]
    section.orientation = WD_ORIENTATION.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for sec in doc.sections:
        sec.top_margin = Inches(1.0)
        sec.bottom_margin = Inches(1.0)
        sec.left_margin = Inches(1.0)
        sec.right_margin = Inches(1.0)
        sec.header_distance = Inches(0.492)
        sec.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, COLORS["navy"], 18, 10),
        ("Heading 2", 13, COLORS["navy"], 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    bullet = styles["List Bullet"]
    bullet.font.name = "Calibri"
    bullet.font.size = Pt(11)
    bullet.paragraph_format.left_indent = Inches(0.375)
    bullet.paragraph_format.first_line_indent = Inches(-0.188)
    bullet.paragraph_format.space_after = Pt(4)
    bullet.paragraph_format.line_spacing = 1.25

    footer = section.footer.paragraphs[0]
    footer.text = "Revisão textual - Passagem de Comando 5ª Cia PM"
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor.from_string(COLORS["muted"])


def add_cover(doc, soup):
    h1 = soup.select_one(".hero h1")
    title = h1.get_text(" ", strip=True) if h1 else "Passagem de comando da 5ª Cia PM"
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Documento de uso interno")
    r.font.size = Pt(10)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string("7A5A00")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(title)
    r.font.name = "Calibri"
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(COLORS["navy_dark"])

    meta = []
    for item in soup.select(".meta-item"):
        label = item.select_one("span")
        value = item.select_one("strong")
        if label and value:
            meta.append((label.get_text(" ", strip=True), value.get_text(" ", strip=True)))
    if meta:
        table = doc.add_table(rows=0, cols=2)
        table.style = "Table Grid"
        set_table_widths(table, [1900, 7460])
        for idx, (label, value) in enumerate(meta):
            row = table.add_row()
            set_cell_shading(row.cells[0], COLORS["header_fill"])
            if idx % 2:
                set_cell_shading(row.cells[1], COLORS["row_alt"])
            for cell in row.cells:
                set_cell_margins(cell)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            row.cells[0].paragraphs[0].add_run(label).bold = True
            row.cells[1].paragraphs[0].add_run(value).bold = True
        doc.add_paragraph()


def table_widths(headers):
    count = len(headers)
    joined = " ".join(headers).lower()
    if count == 4 and "situação" in joined:
        return [1800, 4700, 1700, 1160]
    if count == 4 and "preparar" in joined:
        return [1400, 3300, 1700, 2960]
    if count == 4:
        return [1500, 3600, 1800, 2460]
    if count == 3 and "ok" in joined:
        return [2300, 6100, 960]
    if count == 3:
        return [2300, 3200, 3860]
    if count == 2:
        return [2200, 7160]
    return [int(9360 / count)] * count


def add_table_from_html(doc, table_tag):
    rows = table_tag.select("tr")
    if not rows:
        return
    col_count = max(len(r.find_all(["th", "td"], recursive=False)) for r in rows)
    table = doc.add_table(rows=0, cols=col_count)
    table.style = "Table Grid"
    headers = [c.get_text(" ", strip=True) for c in rows[0].find_all(["th", "td"], recursive=False)]
    set_table_widths(table, table_widths(headers))

    for ridx, tr in enumerate(rows):
        docx_row = table.add_row()
        set_row_cant_split(docx_row)
        cells = tr.find_all(["th", "td"], recursive=False)
        for cidx, cell_tag in enumerate(cells):
            cell = docx_row.cells[cidx]
            set_cell_margins(cell, top=90, bottom=90, start=130, end=130)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if ridx == 0:
                set_cell_shading(cell, COLORS["navy"])
            elif ridx % 2 == 0:
                set_cell_shading(cell, COLORS["row_alt"])
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.12
            add_runs(p, cell_tag)
            collapsed = " ".join(p.text.split())
            for run in p.runs:
                run.text = ""
            if collapsed:
                if p.runs:
                    p.runs[0].text = collapsed
                else:
                    p.add_run(collapsed)
            for run in p.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(8.6 if ridx else 8.2)
                if ridx == 0:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor.from_string("FFFFFF")
            if cidx >= max(0, len(cells) - 2) and ridx > 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_repeat_table_header(table.rows[0])
    doc.add_paragraph()


def render_block(doc, tag):
    if not isinstance(tag, Tag):
        return
    if tag.name == "h3":
        doc.add_paragraph(tag.get_text(" ", strip=True), style="Heading 3")
    elif tag.name == "p":
        add_html_paragraph(doc, tag)
    elif tag.name in {"ul", "ol"}:
        for li in tag.find_all("li", recursive=False):
            add_html_paragraph(doc, li, bullet=True)
    elif tag.name == "table":
        add_table_from_html(doc, tag)
    elif tag.name == "aside":
        label = tag.find("strong")
        if label:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(label.get_text(" ", strip=True))
            run.bold = True
            run.font.color.rgb = RGBColor.from_string("9B1C1C")
        for li in tag.find_all("li"):
            add_html_paragraph(doc, li, bullet=True)
    elif tag.name in {"div"}:
        direct_table = tag.find("table", recursive=False)
        if direct_table:
            add_table_from_html(doc, direct_table)
        else:
            for child in tag.children:
                render_block(doc, child)


def build():
    soup = BeautifulSoup(HTML.read_text(encoding="utf-8"), "html.parser")
    doc = Document()
    style_document(doc)
    add_cover(doc, soup)

    for details in soup.select("details.report-section"):
        summary = details.select_one("summary span")
        if not summary:
            continue
        doc.add_paragraph(summary.get_text(" ", strip=True), style="Heading 1")
        content = details.select_one(".section-content")
        if not content:
            continue
        for child in content.children:
            render_block(doc, child)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
